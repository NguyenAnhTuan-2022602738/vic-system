import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = docx.Document()
    
    # Tiêu đề
    title = doc.add_heading('CƠ SỞ LÝ THUYẾT VÀ CÔNG THỨC TÍCH HỢP CẢM XÚC TIN TỨC VÀO DỰ BÁO', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Nền tảng
    doc.add_heading('1. Nền tảng lý thuyết (Theoretical Framework)', level=1)
    doc.add_paragraph(
        'Phương pháp tiếp cận mô hình tác động (Impact Model) trong hệ thống dự báo được xây dựng dựa trên sự giao thoa giữa Tài chính Hành vi (Behavioral Finance) và Mô hình Định lượng (Quantitative Finance). Việc đưa yếu tố cảm xúc (sentiment) vào điều chỉnh lợi nhuận kỳ vọng và rủi ro được lấy cảm hứng từ các nền tảng sau:'
    )
    
    ul1 = doc.add_paragraph(style='List Bullet')
    ul1.add_run('Mô hình Black-Litterman (1992): ').bold = True
    ul1.add_run('Cho phép nhà đầu tư kết hợp "góc nhìn" (views) vào dự báo định lượng cơ sở. Trong hệ thống này, "góc nhìn" được định lượng hóa tự động bằng kết quả phân tích cảm xúc từ LLM (Large Language Model).')
    
    ul2 = doc.add_paragraph(style='List Bullet')
    ul2.add_run('Lý thuyết Biến động và Cú sốc Thông tin (ARCH/GARCH - Engle, 1982): ').bold = True
    ul2.add_run('Các luồng thông tin mới (dù là tin tốt hay tin xấu) đều phá vỡ trạng thái cân bằng và làm tăng độ biến động (volatility) của thị trường. Điều này lý giải tại sao hệ thống sử dụng giá trị tuyệt đối của cảm xúc (|sentiment|) để điều chỉnh nới rộng độ rủi ro.')
    
    ul3 = doc.add_paragraph(style='List Bullet')
    ul3.add_run('Nghiên cứu của Tetlock (2007) - "Giving Content to Investor Sentiment": ').bold = True
    ul3.add_run('Chứng minh bằng thực nghiệm rằng việc định lượng ngôn ngữ học từ các bài báo truyền thông tài chính có sức mạnh dự báo đáng kể đến xu hướng giá và khối lượng giao dịch trong ngắn hạn.')

    # 2. Công thức Lợi nhuận
    doc.add_heading('2. Công thức Điều chỉnh Lợi nhuận Kỳ vọng (μ_adj)', level=1)
    
    formula1 = doc.add_paragraph('μ_adj = μ × (1 + α × sentiment × impact)')
    formula1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula1.runs[0].font.size = Pt(14)
    formula1.runs[0].bold = True
    
    doc.add_paragraph('Giải thích các biến số:').bold = True
    
    p_mu = doc.add_paragraph(style='List Bullet')
    p_mu.add_run('μ (Base Expected Return): ').bold = True
    p_mu.add_run('Lợi nhuận kỳ vọng cơ sở dự báo thuần túy từ chuỗi thời gian của giá (ví dụ: mô hình LSTM).')
    
    p_sen = doc.add_paragraph(style='List Bullet')
    p_sen.add_run('sentiment ∈ [-1, 1]: ').bold = True
    p_sen.add_run('Điểm phân cực cảm xúc trích xuất từ văn bản tin tức thông qua LLM. Giá trị > 0 biểu thị sắc thái tích cực kéo theo sự lạc quan mua vào, giá trị < 0 là tiêu cực tạo áp lực bán.')
    
    p_imp = doc.add_paragraph(style='List Bullet')
    p_imp.add_run('impact ∈ [0, 1]: ').bold = True
    p_imp.add_run('Trọng số tác động đánh giá mức độ nghiêm trọng/tầm ảnh hưởng của sự kiện trong bản tin.')
    
    p_alpha = doc.add_paragraph(style='List Bullet')
    p_alpha.add_run('α (Alpha - Sentiment Weight): ').bold = True
    p_alpha.add_run('Tham số hệ số nhạy cảm. Được dùng để hiệu chỉnh tỷ lệ ảnh hưởng của tin tức lên giá (tùy theo tính chất ngành hoặc cổ phiếu).')
    
    doc.add_paragraph('Cơ sở học thuật: Công thức này tuân theo dạng điều chỉnh tuyến tính (Linear Adjustment Form) nhằm lồng ghép trực tiếp tác động lan truyền của cảm xúc thị trường (Market Sentiment Spread) vào định giá sinh lời trong ngắn hạn, một phương pháp phổ biến trong khai phá dữ liệu tài chính (Textual Sentiment Analysis).')

    # 3. Công thức rủi ro
    doc.add_heading('3. Công thức Điều chỉnh Rủi ro / Độ bất định (σ_adj)', level=1)
    
    formula2 = doc.add_paragraph('σ_adj = σ × (1 + β × |sentiment| × impact)')
    formula2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula2.runs[0].font.size = Pt(14)
    formula2.runs[0].bold = True
    
    doc.add_paragraph('Giải thích các biến số:').bold = True
    
    p_sigma = doc.add_paragraph(style='List Bullet')
    p_sigma.add_run('σ (Base Volatility): ').bold = True
    p_sigma.add_run('Mức độ rủi ro (độ lệch chuẩn/dải bất định) cơ sở của cổ phiếu theo điều kiện thông thường.')
    
    p_abs_sen = doc.add_paragraph(style='List Bullet')
    p_abs_sen.add_run('|sentiment|: ').bold = True
    p_abs_sen.add_run('Giá trị tuyệt đối của điểm cảm xúc. Biểu thị cường độ của cảm xúc mà không màng đến hướng đi.')
    
    p_beta = doc.add_paragraph(style='List Bullet')
    p_beta.add_run('β (Beta - Uncertainty Weight): ').bold = True
    p_beta.add_run('Hệ số khuếch đại rủi ro từ yếu tố thông tin (Information Leverage Effect).')
    
    doc.add_paragraph('Cơ sở học thuật: Theo nguyên lý của thị trường hiệu quả và mô hình rủi ro có điều kiện (Conditional Volatility Models), bất kỳ thông tin nào đi chệch khỏi kỳ vọng (dù là vĩ mô xuất sắc hay khủng hoảng bất ngờ) đều kích hoạt dòng tiền biến động lớn. Việc sử dụng |sentiment| đảm bảo rằng biến động giá luôn được dự phóng rộng hơn (tăng σ_adj) mỗi khi có tin tức lớn (High Impact).')

    # 4. Tài liệu tham khảo
    doc.add_heading('4. Tài liệu tham khảo (References)', level=1)
    
    ref1 = doc.add_paragraph('[1] Black, F., & Litterman, R. (1992). ')
    ref1.add_run('Global Portfolio Optimization. ').italic = True
    ref1.add_run('Financial Analysts Journal, 48(5), 28-43.')
    
    ref2 = doc.add_paragraph('[2] Engle, R. F. (1982). ')
    ref2.add_run('Autoregressive Conditional Heteroscedasticity with Estimates of the Variance of United Kingdom Inflation. ').italic = True
    ref2.add_run('Econometrica, 50(4), 987-1007.')
    
    ref3 = doc.add_paragraph('[3] Tetlock, P. C. (2007). ')
    ref3.add_run('Giving Content to Investor Sentiment: The Role of Media in the Stock Market. ').italic = True
    ref3.add_run('The Journal of Finance, 62(3), 1139-1168.')
    
    ref4 = doc.add_paragraph('[4] Kearney, C., & Liu, S. (2014). ')
    ref4.add_run('Textual sentiment in finance: A survey of methods and models. ').italic = True
    ref4.add_run('International Review of Financial Analysis, 33, 171-185.')

    # Lưu file
    doc.save('GiaiThich_CongThuc_LLM_Sentiment.docx')
    print("Document saved successfully to GiaiThich_CongThuc_LLM_Sentiment.docx")

if __name__ == '__main__':
    create_document()
